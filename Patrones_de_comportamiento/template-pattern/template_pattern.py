import os
from abc import ABC, abstractmethod
from typing import Any

# Librerías para conversión
from fpdf import FPDF
import markdown2
from docx import Document
from bs4 import BeautifulSoup

class ConvertidorArchivo(ABC):
    def convertir(self, contenido: str, titulo: str) -> None:
        """
        Método template que define la estructura general del proceso de conversión.
        """
        self.validar_contenido(contenido)
        self.crear_encabezado(titulo)
        self.preprocesar_contenido(contenido)
        self.convertir_contenido(contenido)
        self.guardar_archivo()
        self.mostrar_confirmacion()

    def validar_contenido(self, contenido: str) -> None:
        """Valida que el contenido no esté vacío"""
        if not contenido:
            raise ValueError("El contenido no puede estar vacío")

    def crear_encabezado(self, titulo: str) -> None:
        """Crea un encabezado para el documento"""
        self.titulo = titulo
        self.encabezado = f"# {titulo}\n\n"

    def preprocesar_contenido(self, contenido: str) -> None:
        """
        Preprocesa el contenido para eliminar tags HTML si es necesario
        """
        # Usando BeautifulSoup para limpiar contenido HTML
        soup = BeautifulSoup(contenido, 'html.parser')
        self.texto_plano = soup.get_text(strip=True)

    @abstractmethod
    def convertir_contenido(self, contenido: str) -> None:
        """Método abstracto para convertir contenido"""
        pass

    @abstractmethod
    def guardar_archivo(self) -> None:
        """Método abstracto para guardar archivo"""
        pass

    def mostrar_confirmacion(self) -> None:
        """Muestra mensaje de confirmación de conversión"""
        print(f"✅ Conversión completada para {self.__class__.__name__}")

class ConvertidorPDF(ConvertidorArchivo):
    def convertir_contenido(self, contenido: str) -> None:
        """Convierte contenido a PDF"""
        self.pdf = FPDF()
        self.pdf.add_page()
        
        # Configuración de fuentes
        self.pdf.set_font("Arial", size=12)
        
        # Añadir título
        self.pdf.set_font("Arial", 'B', 16)
        self.pdf.cell(0, 10, txt=self.titulo, ln=True, align='C')
        self.pdf.ln(10)
        
        # Restaurar fuente normal
        self.pdf.set_font("Arial", size=12)
        
        # Añadir contenido
        self.pdf.multi_cell(0, 10, txt=self.texto_plano)

    def guardar_archivo(self) -> None:
        """Guarda el archivo PDF"""
        # Crear directorio si no existe
        output_dir = os.path.join(os.path.dirname(__file__), 'output')
        os.makedirs(output_dir, exist_ok=True)
        
        # Guardar PDF
        output_path = os.path.join(output_dir, f"{self.titulo}.pdf")
        self.pdf.output(output_path)
        print(f"📄 PDF guardado en: {output_path}")

class ConvertidorMarkdown(ConvertidorArchivo):
    def convertir_contenido(self, contenido: str) -> None:
        """Convierte contenido a Markdown"""
        # Convertir texto plano a Markdown
        self.markdown_content = self.encabezado + self.texto_plano

    def guardar_archivo(self) -> None:
        """Guarda el archivo Markdown"""
        # Crear directorio si no existe
        output_dir = os.path.join(os.path.dirname(__file__), 'output')
        os.makedirs(output_dir, exist_ok=True)
        
        # Guardar Markdown
        output_path = os.path.join(output_dir, f"{self.titulo}.md")
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(self.markdown_content)
        print(f"📝 Markdown guardado en: {output_path}")

class ConvertidorWord(ConvertidorArchivo):
    def convertir_contenido(self, contenido: str) -> None:
        """Convierte contenido a Word"""
        # Crear documento Word
        self.documento = Document()
        
        # Añadir título como encabezado
        self.documento.add_heading(self.titulo, level=1)
        
        # Añadir párrafo de contenido
        self.documento.add_paragraph(self.texto_plano)

    def guardar_archivo(self) -> None:
        """Guarda el archivo Word"""
        # Crear directorio si no existe
        output_dir = os.path.join(os.path.dirname(__file__), 'output')
        os.makedirs(output_dir, exist_ok=True)
        
        # Guardar Word
        output_path = os.path.join(output_dir, f"{self.titulo}.docx")
        self.documento.save(output_path)
        print(f"📃 Documento Word guardado en: {output_path}")

def main():
    # Ejemplo de contenido HTML
    contenido_html = """
    <h1>Ejemplo de Conversión</h1>
    <p>Este es un párrafo de ejemplo que será convertido a diferentes formatos.</p>
    <ul>
        <li>PDF</li>
        <li>Markdown</li>
        <li>Word</li>
    </ul>
    """
    
    titulo = "Documento de Ejemplo"

    # Conversores
    conversores = [
        ConvertidorPDF(),
        ConvertidorMarkdown(),
        ConvertidorWord()
    ]

    # Convertir usando cada conversor
    for conversor in conversores:
        print("\n" + "="*50)
        conversor.convertir(contenido_html, titulo)
        print("="*50)

if __name__ == "__main__":
    main()